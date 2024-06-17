import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import StaleElementReferenceException
from docx import Document
from docx.shared import Pt

# Set up Chrome options
chrome_options = Options()
chrome_options.add_argument("--headless")  # Run in headless mode
chrome_options.add_argument("--disable-gpu")  # Disable GPU acceleration

# Set up the WebDriver
driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
wait = WebDriverWait(driver, 10)  # Increase wait time if necessary

# URL of the webpage to scrape
url = 'https://books.toscrape.com/'

# Open the webpage
driver.get(url)

# List to hold book data
books_data = []

# Function to scrape data from a page
def scrape_page():
    products = wait.until(EC.presence_of_all_elements_located((By.CLASS_NAME, 'product_pod')))

    for product in products:
        try:
            # Get book URL to navigate into detail page
            book_url = product.find_element(By.TAG_NAME, 'h3').find_element(By.TAG_NAME, 'a').get_attribute('href')

            # Get book title
            title = product.find_element(By.TAG_NAME, 'h3').find_element(By.TAG_NAME, 'a').get_attribute('title')

            # Navigate to the book detail page
            driver.get(book_url)

            # Get price
            price = driver.find_element(By.CLASS_NAME, 'price_color').text

            # Get stock status
            stock_status = driver.find_element(By.CLASS_NAME, 'instock').text.strip()

            # Get rating
            rating = driver.find_element(By.CLASS_NAME, 'star-rating').get_attribute('class').split()[-1]

            # Get description
            description = driver.find_element(By.ID, 'product_description')
            if description:
                description = description.find_element(By.XPATH, 'following-sibling::p').text
            else:
                description = "No description available"

            # Get category
            category = driver.find_elements(By.CSS_SELECTOR, '.breadcrumb li a')[2].text

            # Get product information
            product_info_elements = driver.find_elements(By.CSS_SELECTOR, 'table.table.table-striped tr')
            product_info = {}
            for elem in product_info_elements:
                key = elem.find_element(By.TAG_NAME, 'th').text
                value = elem.find_element(By.TAG_NAME, 'td').text
                product_info[key] = value

            # Append the data to the list
            books_data.append({
                'title': title,
                'price': price,
                'stock_status': stock_status,
                'rating': rating,
                'description': description,
                'category': category,
                'product_info': product_info
            })

            # Go back to the product list page
            driver.back()

        except StaleElementReferenceException:
            print("StaleElementReferenceException occurred. Retrying...")
            continue

# Loop through the first 5 pages
for page_num in range(1, 6):
    print(f"Scraping page {page_num}...")
    scrape_page()

    # Move to the next page if it exists
    try:
        next_button = wait.until(EC.element_to_be_clickable((By.CLASS_NAME, 'next')))
        next_button.click()
    except:
        break  # Break the loop if there's no next button or if the exception persists

# Close the browser
driver.quit()

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading('Books Data', level=1)

# Add book data to the document
for book in books_data:
    doc.add_heading(book['title'], level=2)
    doc.add_paragraph(f"Price: {book['price']}")
    doc.add_paragraph(f"Stock Status: {book['stock_status']}")
    doc.add_paragraph(f"Rating: {book['rating']}")
    doc.add_paragraph(f"Description: {book['description']}")
    doc.add_paragraph(f"Category: {book['category']}")
    
    # Add product information
    doc.add_heading('Product Information', level=3)
    for key, value in book['product_info'].items():
        doc.add_paragraph(f"{key}: {value}")

    doc.add_paragraph("\n")

# Save the document
doc.save('books_data.docx')

print("Data has been written to books_data.docx")
